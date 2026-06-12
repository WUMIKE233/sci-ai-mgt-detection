"""Audit generated DOCX submission files.

This complements the Markdown readiness audit by reading the generated Word
files directly. It checks for submission-facing defects that can appear only
after conversion, such as leaked Markdown markers or author identifiers in the
double-anonymized manuscript file.
"""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document


OUTPUT_DIR = Path("outputs/readiness")
REPORT_PATH = Path("docs/submission_docx_qa.md")
JSON_PATH = OUTPUT_DIR / "submission_docx_qa.json"

FILES = {
    "anonymous_manuscript": Path("submission_package/manuscript_draft_eswa.docx"),
    "title_page": Path("submission_package/title_page_author_details.docx"),
    "cover_letter": Path("submission_package/cover_letter_draft.docx"),
    "declarations": Path("submission_package/declarations_draft.docx"),
    "highlights": Path("submission_package/highlights.docx"),
    "orcid_information": Path("submission_package/orcid_information.docx"),
    "author_metadata": Path("submission_package/author_metadata_template.docx"),
}


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def check(name: str, status: str, evidence: str, action: str = "") -> dict[str, str]:
    return {"name": name, "status": status, "evidence": evidence, "action": action}


def status(ok: bool, blocker: bool = True) -> str:
    if ok:
        return "pass"
    return "blocker" if blocker else "warn"


def collect_author_identifiers(title_page_text: str) -> list[str]:
    """Derive identifiers from the separate title page without hard-coding PII."""
    identifiers: set[str] = set()

    for email in re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", title_page_text):
        identifiers.add(email)

    for orcid in re.findall(r"\b\d{4}-\d{4}-\d{4}-\d{3}[\dX]\b", title_page_text):
        identifiers.add(orcid)

    for phone in re.findall(r"\+\d[\d\s-]{7,}\d", title_page_text):
        identifiers.add(phone.replace(" ", ""))
        identifiers.add(phone)

    lines = [line.strip() for line in title_page_text.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        if line in {"Author", "Corresponding Author"} and idx + 1 < len(lines):
            name_line = lines[idx + 1]
            identifiers.add(name_line)
            identifiers.update(part.strip() for part in re.split(r"[()]", name_line) if len(part.strip()) >= 2)
        if "university" in line.lower() or "institute" in line.lower():
            identifiers.add(line)
            identifiers.update(part.strip() for part in line.split(",") if len(part.strip()) >= 4)

    return sorted(term for term in identifiers if len(term) >= 2)


def main() -> None:
    checks: list[dict[str, str]] = []
    texts: dict[str, str] = {}

    for label, path in FILES.items():
        exists = path.exists() and path.stat().st_size > 0
        checks.append(check(f"{label} exists", status(exists), f"{path}; exists={exists}; size={path.stat().st_size if path.exists() else 0}"))
        if exists:
            texts[label] = docx_text(path)
        else:
            texts[label] = ""

    manuscript = texts["anonymous_manuscript"]
    title_page = texts["title_page"]
    cover_letter = texts["cover_letter"]
    declarations = texts["declarations"]
    highlights = texts["highlights"]

    author_leak_patterns = collect_author_identifiers(title_page)
    leaks = []
    lower_manuscript = manuscript.lower()
    for identifier in author_leak_patterns:
        normalized = identifier.replace(" ", "")
        if identifier.lower() in lower_manuscript or normalized.lower() in lower_manuscript.replace(" ", ""):
            leaks.append(identifier)
    checks.append(
        check(
            "anonymous manuscript contains no author identifiers",
            status(not leaks),
            f"author_identifier_hits={leaks}",
            "Remove author, affiliation, email, phone, and ORCID from the anonymous manuscript file.",
        )
    )

    placeholder_hits = re.findall(r"RESULT_REQUIRED|\bTODO\b|\bTBD\b|Current preliminary evidence|pending final|not final", manuscript, flags=re.I)
    checks.append(
        check(
            "anonymous manuscript has no placeholders or old proof text",
            status(not placeholder_hits),
            f"placeholder_or_old_text_hits={len(placeholder_hits)}",
            "Regenerate the DOCX from the corrected Markdown sources.",
        )
    )

    numeric_hits = re.findall(r"(?m)^\s*\[\d+\]|\[[0-9]+(?:,\s*[0-9]+)*\]", manuscript)
    checks.append(
        check(
            "anonymous manuscript uses non-numbered APA-style citations",
            status(not numeric_hits),
            f"numeric_reference_hits={len(numeric_hits)}",
            "Use author-year citations and APA-style reference entries.",
        )
    )

    markdown_hits = re.findall(r"(?<!\*)\*[^*\n]{2,}\*(?!\*)|`[^`]+`", manuscript)
    checks.append(
        check(
            "anonymous manuscript has no visible Markdown emphasis markers",
            status(not markdown_hits),
            f"markdown_marker_hits={len(markdown_hits)}",
            "Fix DOCX conversion for italics, bold, or inline-code markers.",
        )
    )

    references_present = "\nReferences\n" in f"\n{manuscript}\n"
    reference_section = manuscript.split("\nReferences\n", 1)[-1] if references_present else ""
    reference_lines = [line for line in reference_section.splitlines() if re.match(r"^[A-Z][A-Za-z' -]+,\s", line)]
    checks.append(
        check(
            "reference section is present and populated",
            status(references_present and len(reference_lines) >= 15),
            f"references_present={references_present}; reference_like_lines={len(reference_lines)}",
            "Ensure the manuscript DOCX includes the complete APA reference list.",
        )
    )

    table_figure_terms = ["Empirical Tables", "Table 1.", "Table 2.", "Table 3.", "Empirical Figures", "Figure 1.", "Figure 2.", "Figure 3."]
    missing_terms = [term for term in table_figure_terms if term not in manuscript]
    checks.append(
        check(
            "tables and figures are included in manuscript DOCX",
            status(not missing_terms),
            f"missing_terms={missing_terms}",
            "Regenerate manuscript assets and rebuild the DOCX.",
        )
    )

    title_author_ready = (
        "Author" in title_page
        and "Corresponding Author" in title_page
        and re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", title_page)
        and re.search(r"\b\d{4}-\d{4}-\d{4}-\d{3}[\dX]\b", title_page)
        and ("University" in title_page or "Institute" in title_page)
    )
    checks.append(
        check(
            "title page contains required author details",
            status(title_author_ready),
            f"title_page_author_details_ready={title_author_ready}",
            "Keep author details on the separate title page, not in the anonymous manuscript.",
        )
    )

    checks.append(
        check(
            "cover letter and declarations contain final scope language",
            status(
                "Paraphrase-Style Distribution Shifts" in cover_letter
                and ("not redistributed" in declarations or "not be redistributed" in declarations)
                and "RAID" in declarations
                and ("GitHub" in declarations or "github.com" in declarations.lower())
            ),
            "cover_title_synced_and_data_scope_present="
            + str(
                "Paraphrase-Style Distribution Shifts" in cover_letter
                and ("not redistributed" in declarations or "not be redistributed" in declarations)
                and "RAID" in declarations
                and ("GitHub" in declarations or "github.com" in declarations.lower())
            ),
            "Rebuild cover letter and declarations from the corrected sources.",
        )
    )

    highlight_lines = [line.strip("- ").strip() for line in highlights.splitlines() if line.strip() and line.strip() != "Highlights"]
    long_highlights = [line for line in highlight_lines if len(line) > 85]
    checks.append(
        check(
            "highlights DOCX satisfies ESWA count and length",
            status(3 <= len(highlight_lines) <= 5 and not long_highlights),
            f"highlight_count={len(highlight_lines)}; long_highlights={len(long_highlights)}",
            "Keep 3-5 highlights and no more than 85 characters each.",
        )
    )

    counts = Counter(item["status"] for item in checks)
    overall = "BLOCKED" if counts["blocker"] else ("WARN" if counts["warn"] else "PASS")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    payload = {"overall_status": overall, "status_counts": dict(counts), "checks": checks}
    JSON_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Submission DOCX QA",
        "",
        f"Overall status: **{overall}**",
        "",
        "| Check | Status | Evidence | Action |",
        "|---|---|---|---|",
    ]
    for item in checks:
        lines.append(
            "| {name} | {status} | {evidence} | {action} |".format(
                name=item["name"],
                status=item["status"],
                evidence=item["evidence"].replace("|", "\\|"),
                action=item["action"].replace("|", "\\|"),
            )
        )
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps({"json": str(JSON_PATH), "report": str(REPORT_PATH), "overall_status": overall}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
