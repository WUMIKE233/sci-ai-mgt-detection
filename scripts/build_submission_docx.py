"""Build editable DOCX drafts for the submission package."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("submission_package")
FIGURE_WIDTH = Inches(6.2)


def set_run_font(run, name: str = "Calibri", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = ""
    props.subject = ""
    props.keywords = ""
    props.comments = ""

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_formatted_text(paragraph, text: str) -> None:
    """Add text with simple Markdown emphasis and inline-code handling."""
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=10)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run)
            run.italic = True
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if re.match(r"^\|\s*-+", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_width(table)
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            text = row[col_idx] if col_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.0)
            add_formatted_text(paragraph, text)
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=(row_idx == 0))
    doc.add_paragraph()


def markdown_to_docx(doc: Document, markdown: str) -> None:
    pending_table: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("|") and line.endswith("|"):
            pending_table.append(line)
            continue
        if pending_table:
            add_markdown_table(doc, pending_table)
            pending_table = []
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.style = doc.styles["Title"]
            set_paragraph_spacing(paragraph, before=0, after=12, line_spacing=1.1)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=18, bold=True)
            run.font.color.rgb = RGBColor.from_string("0B2545")
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)
            add_formatted_text(paragraph, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(paragraph, after=4, line_spacing=1.167)
            add_formatted_text(paragraph, line[2:].strip())
        else:
            paragraph = doc.add_paragraph()
            set_paragraph_spacing(paragraph, after=6, line_spacing=1.1)
            add_formatted_text(paragraph, line)
    if pending_table:
        add_markdown_table(doc, pending_table)


def add_section_break(doc: Document, title: str) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(title, level=1)


def add_picture_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        paragraph = doc.add_paragraph()
        add_formatted_text(paragraph, f"[Missing figure file: {path}]")
        return
    doc.add_picture(str(path), width=FIGURE_WIDTH)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=10, line_spacing=1.0)
    run = paragraph.add_run(caption)
    set_run_font(run, size=9, bold=True)


def build_manuscript() -> Path:
    doc = Document()
    configure_document(doc)
    markdown_to_docx(doc, Path("manuscript/manuscript_draft.md").read_text(encoding="utf-8"))

    add_section_break(doc, "Empirical Tables")
    table1 = Path("manuscript/tables/table1_empirical_metrics_with_ci.md")
    table2 = Path("manuscript/tables/table2_empirical_calibration_selective.md")
    table3 = Path("manuscript/tables/table3_empirical_worst_subgroups.md")
    doc.add_heading("Table 1. Locked Empirical Metrics With Bootstrap Confidence Intervals", level=2)
    markdown_to_docx(doc, table1.read_text(encoding="utf-8") if table1.exists() else "")
    doc.add_heading("Table 2. Calibration and Selective Prediction Metrics", level=2)
    markdown_to_docx(doc, table2.read_text(encoding="utf-8") if table2.exists() else "")
    doc.add_heading("Table 3. Highest-Error Subgroups Under the Locked Protocol", level=2)
    markdown_to_docx(doc, table3.read_text(encoding="utf-8") if table3.exists() else "")

    add_section_break(doc, "Empirical Figures")
    add_picture_if_exists(doc, Path("manuscript/figures/figure1_empirical_accuracy_with_ci.png"), "Figure 1. Locked empirical TF-IDF Logistic Regression accuracy.")
    add_picture_if_exists(doc, Path("manuscript/figures/figure2_empirical_ece.png"), "Figure 2. Calibration error under the locked protocol.")
    add_picture_if_exists(doc, Path("manuscript/figures/figure3_empirical_selective_risk.png"), "Figure 3. Selective prediction risk at fixed coverage.")
    captions = Path("manuscript/figures/empirical_figure_captions.md")
    if captions.exists():
        doc.add_heading("Figure Captions", level=2)
        markdown_to_docx(doc, captions.read_text(encoding="utf-8"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / "manuscript_draft_eswa.docx"
    doc.save(path)
    return path


def build_simple_docx(source: Path, output_name: str) -> Path:
    doc = Document()
    configure_document(doc)
    markdown_to_docx(doc, source.read_text(encoding="utf-8"))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / output_name
    doc.save(path)
    return path


def main() -> None:
    outputs = {
        "manuscript": str(build_manuscript()),
        "title_page": str(build_simple_docx(Path("manuscript/title_page.md"), "title_page_author_details.docx")),
        "cover_letter": str(build_simple_docx(Path("manuscript/cover_letter_draft.md"), "cover_letter_draft.docx")),
        "declarations": str(build_simple_docx(Path("manuscript/declarations.md"), "declarations_draft.docx")),
        "highlights": str(build_simple_docx(Path("manuscript/highlights.md"), "highlights.docx")),
        "orcid_information": str(build_simple_docx(Path("manuscript/orcid_information.md"), "orcid_information.docx")),
        "author_metadata": str(build_simple_docx(Path("docs/author_metadata_template.md"), "author_metadata_template.docx")),
    }
    (OUTPUT_DIR / "submission_docx_manifest.json").write_text(json.dumps(outputs, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(outputs, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
